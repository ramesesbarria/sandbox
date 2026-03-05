from utils.helper import multiTableBlockFinder
from utils.helper import deleteColumnInRow, setTableRowHeight
from utils.helper import iterBlockItems
from utils.helper import removePrefixFromParagraph, move_blocks_after

from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

import re
from state import State
from utils.log import log_message

TABLE_IDENTIFIERS = {
    tuple(['Product', 'Current Value', 'Adjustment to\nValue', 'Proposed Value', 'Indirect Cost Ratio (ICR)', 'Indirect Cost Ratio (ICR)']): {
		"name":"Proposed Investment Portfolio Table",
        "process": 3,
		"processDescription": "Proposed Investment Portfolio",
	},
}

def editIPRTableProperties(document):

	#find all table from TABLE_IDENTIFIERS and save it to table list
	tableList = multiTableBlockFinder(document, TABLE_IDENTIFIERS)
	if tableList:
		updatePIPTableProperties(tableList[0])
		updatePIPParagraphTextColor(tableList[0])
	else:
		State.notRunProgress = State.notRunProgress + "\n3.1 editIPRTableProperties"

def updatePIPTableProperties(table):
	rowCount = 0
	for row in table.rows:
		if len(row.cells) > 1:
			if rowCount == 0:
				deleteColumnInRow(row, 4)
				rowCount = rowCount + 1
			elif rowCount == 1:
				deleteColumnInRow(row, 5)
				deleteColumnInRow(row, 4)
				rowCount = rowCount + 1
			else:
				setTableRowHeight(row, 0.6)
				deleteColumnInRow(row, 5)
				deleteColumnInRow(row, 4)

				cellCount = 0
				for cell in row.cells:
					if cellCount == 0:
						cell.paragraphs[0]  = WD_ALIGN_PARAGRAPH.LEFT
						cellCount = cellCount + 1
					else:
						cell.width = Inches(1.5)
						cell.paragraphs[0]  = WD_ALIGN_PARAGRAPH.RIGHT
	
def updatePIPParagraphTextColor(table):
	numberOfRow = len(table.rows)
	for i in range(numberOfRow):
		row = table.rows[i]
		if len(row.cells) > 1:
			basecolor = ''
			if "Total" in row.cells[0].text:
				baseRow = table.rows[i - 1] 
				baseCell = baseRow.cells[2]
				baseParagraph = baseCell.paragraphs[0]
				baseRun = baseParagraph.runs[0]
				color = baseRun.font.color.rgb
				if (str(color) == "FF0000"):
					cell = row.cells[2]
					for paragraph in cell.paragraphs:
						for run in paragraph.runs:
							run.font.color.rgb = color
	
def removeMDAText(document):
	for block in iterBlockItems(document):
		if isinstance(block, Paragraph):
			blockWithMDA = re.search(r"For MDA Only", block.text, re.IGNORECASE)
			
			if blockWithMDA:
				removePrefixFromParagraph(block, prefix="For MDA Only ")
				break

def moveIMFBSSParagraph(document):

	TABLE_TARGET_IDENTIFIERS = {
		tuple(['', '', 'Upfront', 'Ongoing', 'Ongoing']): {
		"name":"Product cost table",
		"processDescription": "Paste paragraph after product cost table",
		}
	}

	targetBlock = multiTableBlockFinder(document, TABLE_TARGET_IDENTIFIERS)
	blockToMove = []
	inRange = False

	for block in iterBlockItems(document):

		if isinstance(block, Paragraph):
			startBlock = re.search(r"Investment Management Fees", block.text, re.IGNORECASE)
			endBeforeBlock = re.search(r"Your Proposed Asset Allocation", block.text, re.IGNORECASE)
			
			if startBlock:
				inRange = True
			elif endBeforeBlock:
				inRange = False
				break
		if inRange:
			blockToMove.append(block)

	move_blocks_after(blockToMove,targetBlock[0])


def moveSalesAndPurchasesBlock(document):
		
	targetBlock = ''
	blockToMove = []
	inRange = False

	for block in iterBlockItems(document):

		if isinstance(block, Paragraph):
			startBlock = re.search(r"\bSales\b", block.text, re.IGNORECASE)
			endBeforeBlock = re.search(r"\bProposed portfolio(s?)\b", block.text, re.IGNORECASE)
			targetBlock = re.search(r"\bBuy/ Sell Spreads\b", block.text, re.IGNORECASE)
			if startBlock:
				inRange = True
			elif endBeforeBlock:
				inRange = False
			elif targetBlock:
				targetBlock = block
				break

		if inRange:
			blockToMove.append(block)

	move_blocks_after(blockToMove,targetBlock)
		
def moveProposedPortfolioBlock(document):
	targetBlock = ''
	blockToMove = []
	inRange = False

	for block in iterBlockItems(document):

		if isinstance(block, Paragraph):
			startBlock = re.search(r"\bProposed portfolio(s?)\b", block.text, re.IGNORECASE)
			endBeforeBlock = re.search(r"\bCost comparison\b", block.text, re.IGNORECASE)
			targetBlock = re.search(r"\bInvestment Management Fees\b", block.text, re.IGNORECASE)
			if startBlock:
				inRange = True
			elif endBeforeBlock:
				inRange = False
			elif targetBlock:
				targetBlock = block
				break

		if inRange:
			blockToMove.append(block)

	move_blocks_after(blockToMove,targetBlock)


