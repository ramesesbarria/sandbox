from copy import deepcopy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

def add_cell_to_header(row, new_header_text="New Column"):
    """
    Adds a new cell to the header row (row 0) only.
    WARNING: The table will be inconsistent if body rows have fewer cells.
    """
    # Create new cell at the end of header
    # Must use _tc and _tr to append a cell manually


    # Clone last cell
    last_cell = row.cells[-1]
    new_tc = deepcopy(last_cell._tc)  # copy cell structure
    # Clear text in the cloned cell
    for p in new_tc.xpath('.//w:t'):
        p.text = new_header_text

    # Append new cell to the row
    row._tr.append(new_tc)

    return row


def merge_cells_range(row, start_idx, end_idx, font_size):
    """
    Merge  cells in row from start_idx to end_idx (inclusive).
    """

    if start_idx < 0 or end_idx >= len(row.cells) or start_idx >= end_idx:
        raise ValueError("Invalid start or end index for merging header cells")

    first_cell = row.cells[start_idx]
    last_cell = row.cells[end_idx]
        
    first_cell_text = first_cell.text  # already preserved
        
    first_cell.merge(last_cell)

    first_cell.text = ""  # clear merged cell content
    run = first_cell.paragraphs[0].add_run(first_cell_text)
    run.bold = True
    run.font.size = Pt(font_size)

    return row

from docx.enum.text import WD_ALIGN_PARAGRAPH

def align_cell_text(row, cell_idx, alignment=WD_ALIGN_PARAGRAPH.LEFT, text=None):
    """
    Aligns text in a specific cell of a row.
    
    Parameters:
        row: a docx table row object
        cell_idx: index of the cell to align (0-based)
        alignment: WD_ALIGN_PARAGRAPH alignment (LEFT, CENTER, RIGHT, JUSTIFY)
        text: optional string to add if cell is empty
    """
    # Handle negative indices
    if cell_idx < 0:
        cell_idx = len(row.cells) + cell_idx  # e.g., -1 -> last cell

    if cell_idx < 0 or cell_idx >= len(row.cells):
        raise IndexError("Cell index out of range")

    cell = row.cells[cell_idx]

    # Add a paragraph if empty and text is provided
    if not cell.paragraphs and text is not None:
        cell.add_paragraph(text)
    elif text is not None:
        # Replace text of first paragraph
        cell.paragraphs[0].text = text

    # Align all paragraphs in the cell
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
