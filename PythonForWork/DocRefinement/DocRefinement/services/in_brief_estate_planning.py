from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import RGBColor
from utils.helper import iter_block_items
from utils.log import log_message
from state import State

# ── Colour constants ──────────────────────────────────────────────────────────
BLACK = RGBColor(0x00, 0x00, 0x00)


# ── Step 1.2 ─────────────────────────────────────────────────────────────────

def read_estate_planning_entries(document):
    """
    Reads the three Estate Planning Yes/No values from the
    'Your personal and financial position' Personal Details table.

    Returns a dict:
        {
            'will':        'Yes' | 'No' | None,
            'poa':         'Yes' | 'No' | None,
            'guardianship':'Yes' | 'No' | None,
        }
    None means the row was not found.
    """
    entries = {'will': None, 'poa': None, 'guardianship': None}

    ROW_MAP = {
        'do you have a will?':       'will',
        'enduring power of attorney?': 'poa',
        'enduring guardianship?':    'guardianship',
    }

    for table in document.tables:
        for row in table.rows:
            first_cell = row.cells[0].text.strip().lower()
            key = ROW_MAP.get(first_cell)
            if key and len(row.cells) > 1:
                entries[key] = row.cells[1].text.strip()

    return entries


# ── Step 1.3 helpers ──────────────────────────────────────────────────────────

def _build_will_phrase(will_value):
    """'Yes' → 'your Will reviewed'  |  'No' → 'a Will prepared'"""
    if will_value == 'Yes':
        return 'your Will reviewed'
    return 'a Will prepared'


def _build_poa_phrase(poa_value):
    """'Yes' → 'review your'  |  'No' → 'establish an'"""
    if poa_value == 'Yes':
        return 'review your'
    return 'establish an'


def _build_guardianship_phrase(guardianship_value):
    """'Yes' → 'updating your '  |  'No' → 'establishing a '"""
    if guardianship_value == 'Yes':
        return 'updating your '
    return 'establishing a '


def _set_run_black(run):
    run.font.color.rgb = BLACK


def _find_in_brief_estate_table(document):
    """
    Returns the second single-row table in the In Brief section —
    the one containing the estate planning recommendation.
    """
    inside = False
    table_count = 0

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            if block.text.strip() == 'In Brief':
                inside = True
            if inside and block.text.strip() == 'Your personal and financial position':
                break
        elif inside and isinstance(block, Table):
            table_count += 1
            if table_count == 2:
                return block

    return None


def _remove_extra_spacing(document):
    """
    Point 2: Remove the two consecutive empty paragraphs that sit between
    the intro sentence and item 1 in the In Brief section.
    """
    inside = False
    intro_found = False
    removed = 0

    blocks = list(iter_block_items(document))

    for i, block in enumerate(blocks):
        if isinstance(block, Paragraph):
            if block.text.strip() == 'In Brief':
                inside = True
                continue

            if inside and block.text.strip() == 'Your personal and financial position':
                break

            if inside and 'We recommend you implement the following strategies' in block.text:
                intro_found = True
                continue

            # Remove empty paragraphs immediately after the intro sentence,
            # before the first table
            if inside and intro_found and block.text.strip() == '':
                # Check the next non-empty item - if it's a table, these are extra
                p = block._element
                p.getparent().remove(p)
                removed += 1

        elif inside and isinstance(block, Table):
            intro_found = False  # stop removing once we hit a table

    log_message(f"In Brief: removed {removed} extra spacing paragraph(s).")


# ── Step 1.3 main ─────────────────────────────────────────────────────────────

def update_in_brief_estate_planning(document, entries, testamentary_trust_yes):
    """
    Point 2: Remove extra spacing.
    Point 3: Update recommendation text based on Yes/No entries.
    Point 4: Remove or keep the testamentary trust paragraph.
    Point 5: Set all text in the cell to black.
    """

    # Point 2
    _remove_extra_spacing(document)

    # Find the In Brief estate planning table
    table = _find_in_brief_estate_table(document)
    if table is None:
        log_message("In Brief: Could not find estate planning recommendation table.")
        State.notRunProgress += '\n' + State.docname + ' 1.3. In Brief - Estate Planning table not found'
        return

    cell = table.rows[0].cells[2]
    paragraphs = cell.paragraphs

    # Paragraph index 1 = main recommendation sentence
    # Paragraph index 2 = testamentary trust sentence
    if len(paragraphs) < 2:
        log_message("In Brief: Unexpected paragraph count in estate planning cell.")
        State.notRunProgress += '\n' + State.docname + ' 1.3. In Brief - Unexpected cell structure'
        return

    main_para = paragraphs[1]
    runs = main_para.runs

    # ── Point 3: Update the three red placeholder phrases ────────────────────
    # run[1] = 'your Will reviewed / prepared'
    # run[4] = 'establish an / review your'
    # run[6] = 'establishing a / updating your '

    if len(runs) >= 8:
        runs[1].text = _build_will_phrase(entries.get('will'))
        runs[4].text = _build_poa_phrase(entries.get('poa'))
        runs[6].text = _build_guardianship_phrase(entries.get('guardianship'))
    else:
        log_message("In Brief: Unexpected run count in main paragraph — skipping phrase update.")
        State.notRunProgress += '\n' + State.docname + ' 1.3. In Brief - Could not update recommendation phrases'

    # ── Point 4: Testamentary trust paragraph (index 2) ─────────────────────
    if len(paragraphs) >= 3:
        test_para = paragraphs[2]
        if 'testamentary trust' in test_para.text.lower():
            if testamentary_trust_yes:
                # Keep it — just make sure text turns black below
                pass
            else:
                # Remove it
                p = test_para._element
                p.getparent().remove(p)

    # ── Point 5: Set all remaining text in cell to black ────────────────────
    for para in cell.paragraphs:
        for run in para.runs:
            _set_run_black(run)

    log_message(
        f"In Brief: Updated estate planning. "
        f"Will={entries.get('will')}, POA={entries.get('poa')}, "
        f"Guardianship={entries.get('guardianship')}, "
        f"Testamentary={testamentary_trust_yes}"
    )


# ── Entry point ───────────────────────────────────────────────────────────────

def process_in_brief_estate_planning(document, testamentary_trust_yes):
    """
    Main entry point called from document_service.py.

    Args:
        document: the open SOA docx Document object
        testamentary_trust_yes (bool): True if Paraplanning form says Yes
    """
    entries = read_estate_planning_entries(document)

    missing = [k for k, v in entries.items() if v is None]
    if missing:
        log_message(f"In Brief: Could not find estate planning entries: {missing}")
        State.notRunProgress += '\n' + State.docname + f' 1.2. Estate Planning entries not found: {missing}'

    update_in_brief_estate_planning(document, entries, testamentary_trust_yes)