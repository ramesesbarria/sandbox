from utils.helper import table_checker_by_custom_header
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table
import re
from utils.helper import iter_block_items
from utils.helper import set_table_row_height
from state import State

TABLE_IDENTIFIERS = {
    tuple(["Key Assumptions"]): {
		"name": "Key Assumptions Table",
        "process": "4.1.2",
		"processDescription": "Your personal and financial position",
	}
}


def modify_key_assumptions_section(document):
    HEADER_ROWS = (0, 1)
    subHeadingFinderCounter = 0

    modify_header = False
    modify_table = False

    for block in iter_block_items(document):

        # ---------- HEADER ----------
        if isinstance(block, Paragraph):
            if re.search(r"\bKey Assumptions\b", block.text, re.IGNORECASE):
                # prevent double-append
                if "Base Scenario" not in block.text and "Recommended Scenario" not in block.text:
                    suffix = (
                        " - Base Scenario"
                        if subHeadingFinderCounter == 0
                        else " - Recommended Scenario"
                    )

                    new_run = block.add_run(suffix)
                    new_run.bold = True

                    modify_header = True
                    subHeadingFinderCounter += 1
        # ---------- TABLE ----------
        elif  isinstance(block, Table):
            table = table_checker_by_custom_header(block,TABLE_IDENTIFIERS,HEADER_ROWS)
            if table:
                table_modified = False

                for row in table.rows:
                    set_table_row_height(row, 0.6)

                    for idx_c, cell in enumerate(row.cells):
                        if cell.paragraphs:
                            if idx_c == 0:
                                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                            else:
                                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                    table_modified = True

                if table_modified:
                    modify_table = True
            
    if not modify_header:
        State.notRunProgress = State.notRunProgress + "\n" + State.docname + " 4.1.2. Key Assumptions - Header not modified"
    if not modify_table:
        State.notRunProgress = State.notRunProgress + + "\n" + State.docname + " 4.1.2. Key Assumptions - Table not modified"