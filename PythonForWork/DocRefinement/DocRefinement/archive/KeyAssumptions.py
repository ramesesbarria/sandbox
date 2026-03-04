from utils.helper import table_checker_by_custom_header
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table
import re
from utils.helper import iter_block_items
from utils.helper import set_table_row_height

TABLE_IDENTIFIERS = {
    tuple(["Key Assumptions"]): {
		"name": "Key Assumptions Table",
        "process": "4.1.2",
		"processDescription": "Your personal and financial position",
	}
}


def modify_key_assumptions_section(document):
    HEADER_ROWS = (0, 1)
    subHeadingFinderCounter = 0

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            subHeading = re.search(r"Key Assumptions", block.text, re.IGNORECASE)
            if subHeading:
                if subHeadingFinderCounter ==0:
                    new_run = block.add_run(" - Base Scenario")
                    new_run.bold = True
                    subHeadingFinderCounter += 1
                else:
                    new_run = block.add_run(" - Recommended Scenario")
                    new_run.bold = True

        if isinstance(block, Table):
            table = table_checker_by_custom_header(block,TABLE_IDENTIFIERS,HEADER_ROWS)
            if table:
                for row in table.rows:
                    set_table_row_height(row, 0.6)
                    cellCount = 0
                    for cell in row.cells:
                        if cellCount == 0:
                            cell.paragraphs[0]  = WD_ALIGN_PARAGRAPH.LEFT
                            cellCount = cellCount + 1
                        else:
                            cell.paragraphs[0]  = WD_ALIGN_PARAGRAPH.RIGHT

