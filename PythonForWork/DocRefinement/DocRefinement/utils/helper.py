from config import DOC_PATH, UPDATED_PATH
from utils.log import log_message

from docx import Document
from docx.text.paragraph import Paragraph
import re
from docx.table import Table

from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from copy import deepcopy
import os

def open_doc(docName):

    try:
        file_path = DOC_PATH + docName
        print(file_path)
        document = Document(file_path)
        return document

    except Exception as e:
        log_message("File Open: " + str(e))
        return None

def save_doc(document, docName):
    try:
        file_path = UPDATED_PATH + docName
        document.save(file_path)
        return file_path
    except Exception as e:
        log_message("File Saving: " + str(e))
    

def iter_block_items(parent):
    """Yield paragraphs and tables in document order."""
    parent_elm = parent.element.body

    for child in parent_elm:
        if child.tag.endswith('tbl'):
            yield Table(child, parent)
        elif child.tag.endswith('p'):
            yield Paragraph(child, parent)
            
def multi_table_block_finder(document, TABLE_IDENTIFIERS):
    
    found_tables = []
    identifiers_found = []
    for block in iter_block_items(document):
        if isinstance(block, Table):
            headerCells = tuple(cell.text.strip() for cell in block.rows[0].cells)
            table_name = TABLE_IDENTIFIERS.get(headerCells)

            if table_name:
                found_tables.append(block)
                identifiers_found.append(table_name)

    # Log missing tables
    expected_names = [info['name'] for info in TABLE_IDENTIFIERS.values()]
    found_names = [info['name'] for info in identifiers_found]
    missing_tables = set(expected_names) - set(found_names)

    if missing_tables:
        log_message(f"Warning: The following tables were not found: {','.join(missing_tables)}")
    return found_tables

def multi_table_block_finder_with_tag(document, TABLE_IDENTIFIERS):
    
    tableList = []
    for block in iter_block_items(document):
        if isinstance(block, Table):
            headerCells = [cell.text.strip() for cell in block.rows[0].cells]
            header_t = tuple(headerCells)
            table_name = TABLE_IDENTIFIERS.get(header_t)

            if table_name:
                tableList.append([block,header_t[-1]])
    
    return tableList

def find_tables_by_top_leftcells(block, TABLE_IDENTIFIERS):

    if len(block.rows) < 2:
        return None

    # Get first cell of first row
    first_row_first_cell = block.rows[0].cells[0].text.strip()
    # Get first cell of second row
    second_row_first_cell = block.rows[1].cells[0].text.strip()

    header_t = (first_row_first_cell, second_row_first_cell)
    table_meta = TABLE_IDENTIFIERS.get(header_t)
    if table_meta:
        return block

def table_checker_by_custom_header(block, TABLE_IDENTIFIERS, HEADER_ROWS):
    for row_idx in HEADER_ROWS:
        if len(block.rows) > row_idx:
            headerCells = [
                cell.text.replace("\n", " ").strip()
                for cell in block.rows[row_idx].cells
                if cell.text.strip()
            ]

            header_t = tuple(headerCells)
            table_info = TABLE_IDENTIFIERS.get(header_t)

            if table_info:
                return block



def set_table_full_width(table):
    """
    Adjust table to span the full page width, respecting the section margins.
    """
    # Get the section the table belongs to
    # Most reliable: use the first paragraph in the table
    first_paragraph = table.rows[0].cells[0].paragraphs[0]
    section = first_paragraph.part.section_start or first_paragraph._element.getparent().section

    # Or more commonly:
    section = table._parent
    while not hasattr(section, "page_width"):
        section = getattr(section, "_parent", None)
        if section is None:
            # Cannot find section, skip width adjustment
            return

    # Available width = page width minus margins
    available_width = section.page_width - section.left_margin - section.right_margin

    # Set each column width evenly
    num_cols = len(table.columns)
    for column in table.columns:
        column.width = available_width // num_cols

def get_table_width_cm(table, default_width_cm=17):
    """
    Returns table width in cm. If not set, returns default width.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr

    # Check if tblW exists
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None:
        return int(tblW.get(qn('w:w'))) / 567  # convert twips to cm
    else:
        return default_width_cm

def add_cell_to_header(row, new_header_text="New Column"):
    """
    Adds a new cell to the header row (row 0) only.
    WARNING: The table will be inconsistent if body rows have fewer cells.
    """
    # Create new cell at the end of header
    # Must use _tc and _tr to append a cell manually


    # Clone last cell
    last_cell = row.cells[-1]
    new_tc = deepcopy(last_cell._tc)  # copy cell structure
    # Clear text in the cloned cell
    for p in new_tc.xpath('.//w:t'):
        p.text = new_header_text

    # Append new cell to the row
    row._tr.append(new_tc)

    return row

def merge_cells_range(row, start_idx, end_idx, font_size):
    """
    Merge  cells in row from start_idx to end_idx (inclusive).
    """

    if start_idx < 0 or end_idx >= len(row.cells) or start_idx >= end_idx:
        raise ValueError("Invalid start or end index for merging header cells")

    first_cell = row.cells[start_idx]
    last_cell = row.cells[end_idx]
        
    first_cell_text = first_cell.text  # already preserved
        
    first_cell.merge(last_cell)

    first_cell.text = ""  # clear merged cell content
    run = first_cell.paragraphs[0].add_run(first_cell_text)
    run.bold = True
    run.font.size = Pt(font_size)

    return row

def align_cell_text(row, cell_idx, alignment, text=None):
    """
    Aligns text in a specific cell of a row.
    
    Parameters:
        row: a docx table row object
        cell_idx: index of the cell to align (0-based)
        alignment: WD_ALIGN_PARAGRAPH alignment (LEFT, CENTER, RIGHT, JUSTIFY)
        text: optional string to add if cell is empty
    """
    # Handle negative indices
    if cell_idx < 0:
        cell_idx = len(row.cells) + cell_idx  # e.g., -1 -> last cell

    if cell_idx < 0 or cell_idx >= len(row.cells):
        raise IndexError("Cell index out of range")

    cell = row.cells[cell_idx]

    # Add a paragraph if empty and text is provided
    if not cell.paragraphs and text is not None:
        cell.add_paragraph(text)
    elif text is not None:
        # Replace text of first paragraph
        cell.paragraphs[0].text = text

    # Align all paragraphs in the cell
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment

def delete_column_in_row(row, col_idx):
    cell = row.cells[col_idx]
    cell._tc.getparent().remove(cell._tc)


def get_table_row_height_cm(row):

    tr = row._tr
    trPr = tr.trPr
    if trPr is not None:
        trHeight = trPr.find(qn('w:trHeight'))
        if trHeight is not None:
            height = trHeight.get(qn('w:val'))
            heightTwips = int(height)
            return twips_to_cm(heightTwips)
    return None
    
def set_table_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()

    # Remove old trHeight
    for existing in trPr.findall(qn('w:trHeight')):
        trPr.remove(existing)

    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(cm_to_twips(height_cm)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def cm_to_twips(cm):
    return int(cm * 567)

def twips_to_cm(twips):
    return round(float(twips / 567,),2)

def inches_to_twips(inches):
    """Word measures widths in twips (1/1440 inch)."""
    return int(inches * 1440)


def set_paragraph_font(paragraph, fontSize):
    for run in paragraph.runs:
            run.font.size = Pt(fontSize)

def set_row_paragraph_font(row, fontSize):
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(fontSize)

def remove_prefix_from_paragraph(paragraph, prefix):
    text = paragraph.text
    # Use regex to remove the prefix at the start
    new_text = re.sub(rf"^\s*{prefix}\s*", "", text, flags=re.IGNORECASE)
    paragraph.text = new_text

def get_xml_element(obj):
    """Return the underlying XML element whether the object is a paragraph or table."""
    if hasattr(obj, "_p"):
        return obj._p   # paragraph XML
    elif hasattr(obj, "_tbl"):
        return obj._tbl # table XML
    else:
        raise TypeError("Object must be a paragraph or a table.")

def move_blocks_after(blocks_to_move, target_block):
    """
    Move paragraphs/tables so that they appear AFTER the target paragraph/table.

    Parameters:
        blocks_to_move : list of paragraph or table objects
        target_block   : paragraph or table object after which blocks should be inserted
    """
    # XML of target
    target_elem = get_xml_element(target_block)
    parent = target_elem.getparent()

    # XML elements of blocks to move
    move_elems = [get_xml_element(b) for b in blocks_to_move]

    if not move_elems:
        return False
    # Insert index right after the target
    insert_index = parent.index(target_elem) + 1

    # --- INSERT COPIES IN NEW LOCATION ---
    for elem in move_elems:
        new_elem = deepcopy(elem)
        parent.insert(insert_index, new_elem)
        insert_index += 1

    # --- REMOVE ORIGINALS ---
    for elem in move_elems:
        old_parent = elem.getparent()
        if old_parent is not None:
            old_parent.remove(elem)

    return True

def is_row_blank(row):
    """Return True if all cells in a row are empty or whitespace."""
    for cell in row.cells:
        if cell.text.strip():  # If ANY text exists
            return False
    return True

def delete_row(table, row):
    """Remove a row from a table (python-docx internal XML delete)."""
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

def insert_section_break(paragraph):
    p = paragraph._p
    sep = OxmlElement("w:pPr")
    sectPr = OxmlElement("w:sectPr")

    # type="nextPage" => Section Break (Next Page)
    type_tag = OxmlElement("w:type")
    type_tag.set(qn("w:val"), "nextPage")
    sectPr.append(type_tag)

    # --- Set LANDSCAPE ORIENTATION ---
    pgSz = OxmlElement("w:pgSz")
    pgSz.set(qn("w:h"), "15840")   # inches  * 1440
    pgSz.set(qn("w:w"), "11900")   # inches * 1440
    pgSz.set(qn("w:orient"), "portrait")
    sectPr.append(pgSz)

    sep.append(sectPr)

    p.append(sep)

    return sectPr

def insert_landscape_section_before(paragraph):
    """
    Inserts a section break (next page) before a paragraph
    and sets the new section to LANDSCAPE orientation.
    """

    # Get the XML <w:p> of the paragraph
    p = paragraph._p

    # Ensure paragraph has <w:pPr>
    pPr = p.get_or_add_pPr()

    # Create a new <w:sectPr> (section properties)
    sectPr = OxmlElement("w:sectPr")

    # --- Make this a NEXT PAGE section break ---
    type_elm = OxmlElement("w:type")
    type_elm.set(qn("w:val"), "nextPage")
    sectPr.append(type_elm)

    # --- Set LANDSCAPE ORIENTATION ---
    pgSz = OxmlElement("w:pgSz")
    pgSz.set(qn("w:w"), "15840")   # 11 inches  * 1440
    pgSz.set(qn("w:h"), "11900")   # 8.5 inches * 1440
    pgSz.set(qn("w:orient"), "landscape")
    sectPr.append(pgSz)

    # Optional margins (Word defaults if omitted)
    pgMar = OxmlElement("w:pgMar")
    pgMar.set(qn("w:top"), "1134")    
    pgMar.set(qn("w:right"), "1134")
    pgMar.set(qn("w:bottom"), "1134")
    pgMar.set(qn("w:left"), "1134")
    sectPr.append(pgMar)

    # Insert new section break before paragraph
    pPr.append(sectPr)

    return sectPr

def add_breaks_in_between_paragraph(document, START_MARKER, END_MARKER):

    start_block = None
    end_block = None
    inside = False

    # STEP 1 — Find START and END paragraphs
    for p in document.paragraphs:
        
        text = p.text.strip()
        if text == START_MARKER and not inside:
            start_block = p
            inside = True
            continue

        elif inside and text == END_MARKER:
            end_block = p
            break
        elif inside:
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            
    if not start_block or not end_block:
        print("Markers not found.")
        return False
    
    # ---------------------------------------------------------------------------
    # STEP 2 — Insert a SECTION BREAK BEFORE the start marker
    # ---------------------------------------------------------------------------
    # Insert an empty paragraph before start marker
    before_start = start_block.insert_paragraph_before()
    insert_section_break(before_start)
    
    # ---------------------------------------------------------------------------
    # STEP 3 — Insert a SECTION BREAK AFTER the end marker
    # ---------------------------------------------------------------------------
    before_end = end_block.insert_paragraph_before()
    insert_landscape_section_before(before_end)

    return True

def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None
    return paragraph._p