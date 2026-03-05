from utils.helper import find_tables_by_top_leftcells
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table
import re
from utils.helper import iter_block_items 
from utils.helper import set_table_row_height
from utils.helper import add_breaks_in_between_paragraph
from state import State

TABLE_IDENTIFIERS = {
    tuple(["Date","Age"]): {
		"name": "Consolidated CashFlow and Assets Table",
        "process": "4.1.3",
		"processDescription": "Your personal and financial position",
	}
}



def modify_cash_flow_section(document):
    modify_header = False
    modify_table = False

    subHeadingFinderCounter = 0
    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            consolidatedCashflowSubHeading = re.search(r"Consolidated Cashflow", block.text, re.IGNORECASE)
            AssetsTotalSubHeading = re.search(r"Assets - Total - Detail", block.text, re.IGNORECASE)
            if consolidatedCashflowSubHeading or AssetsTotalSubHeading:
                  if "Base Scenario" not in block.text and "Recommended Scenario" not in block.text:
                    suffix = (
                        " - Base Scenario"
                        if subHeadingFinderCounter == 0
                        else " - Recommended Scenario"
                    )

                    new_run = block.add_run(suffix)
                    new_run.bold = True

                    modify_header = True
            
        if isinstance(block, Table):
            table = find_tables_by_top_leftcells(block,TABLE_IDENTIFIERS)
            

            if table:
                table_modified = False
                for row in table.rows:
                    set_table_row_height(row, 0.6)
                    cellCount = 0
                    for cell in row.cells:
                        if cellCount == 0:
                            cell.paragraphs[0]  = WD_ALIGN_PARAGRAPH.LEFT
                            cellCount = cellCount + 1
                        else:
                            cell.paragraphs[0]  = WD_ALIGN_PARAGRAPH.RIGHT
                    table_modified = True
    
                if table_modified:
                    modify_table = True

    if not modify_header:
        State.notRunProgress = State.notRunProgress + "\n" + State.docname + " 4.1.3. Consolidated Cashflow - Header"
    if not modify_table:
        State.notRunProgress = State.notRunProgress + "\n" + State.docname + " 4.1.3. Consolidated Cashflow - Table"

def modify_consolidated_cashflow_section(document):
    START_MARKER = "Consolidated Cashflow"
    END_MARKER   = "Appendix: Investment summaries"


    if not add_breaks_in_between_paragraph(document,START_MARKER,END_MARKER):
        State.notRunProgress = State.notRunProgress + "\n" + State.docname + " 4.1.3. Consolidated Cashflow - Break and Page layout"
    modify_cash_flow_section(document)