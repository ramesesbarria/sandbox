from utils.helper import find_tables_by_top_leftcells
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table
import re
from utils.helper import iter_block_items 
from utils.helper import set_table_row_height
from utils.helper import add_breaks_in_between_paragraph

TABLE_IDENTIFIERS = {
    tuple(["Date","Age"]): {
		"name": "Consolidated CashFlow and Assets Table",
        "process": "4.1.3",
		"processDescription": "Your personal and financial position",
	}
}
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_BREAK

def is_paragraph_empty_or_break(paragraph):
    return not paragraph.text.strip()

def updateCashFlowSection(document):
    subHeadingFinderCounter = 0
    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            consolidatedCashflowSubHeading = re.search(r"Consolidated Cashflow", block.text, re.IGNORECASE)
            AssetsTotalSubHeading = re.search(r"Assets - Total - Detail", block.text, re.IGNORECASE)
            if consolidatedCashflowSubHeading or AssetsTotalSubHeading:
                if subHeadingFinderCounter ==0:
                    new_run = block.add_run(" - Base Scenario")
                    new_run.bold = True
                    subHeadingFinderCounter += 1
                else:
                    new_run = block.add_run(" - Recommended Scenario")
                    new_run.bold = True

            if AssetsTotalSubHeading:
                AssetsTotalSubHeading = block
            
        if isinstance(block, Table):
            table = find_tables_by_top_leftcells(block,TABLE_IDENTIFIERS)
            if table:
                for row in table.rows:
                    set_table_row_height(row, 0.6)
                    cellCount = 0
                    for cell in row.cells:
                        if cellCount == 0:
                            cell.paragraphs[0]  = WD_ALIGN_PARAGRAPH.LEFT
                            cellCount = cellCount + 1
                        else:
                            cell.paragraphs[0]  = WD_ALIGN_PARAGRAPH.RIGHT

    return document


def modify_consolidated_cashflow_section(document):
    START_MARKER = "Consolidated Cashflow"
    END_MARKER   = "Appendix: Investment summaries"
    add_breaks_in_between_paragraph(document,START_MARKER,END_MARKER)
    updateCashFlowSection(document)