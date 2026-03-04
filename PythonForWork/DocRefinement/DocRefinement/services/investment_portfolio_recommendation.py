from utils.helper import multi_table_block_finder
from utils.helper import delete_column_in_row, set_table_row_height
from utils.helper import iter_block_items
from utils.helper import remove_prefix_from_paragraph, move_blocks_after
from utils.helper import remove_paragraph

from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

import re
from state import State
from utils.log import log_message

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import copy

TABLE_IDENTIFIERS_ALL = {
    tuple(['Product', 'Current Value', 'Adjustment to\nValue', 'Proposed Value', 'Indirect Cost Ratio (ICR)']): {
		"name":"Proposed Investment Portfolio Table",
        "process": 3,
		"processDescription": "Proposed Investment Portfolio",
	},
	tuple(['Product', 'Current Value', 'Adjustment to\nValue', 'Proposed Value', 'Indirect Cost Ratio (ICR)', 'Indirect Cost Ratio (ICR)']): {
		"name":"Proposed Investment Portfolio Table",
        "process": 3,
		"processDescription": "Proposed Investment Portfolio",
	},
	tuple(['Product', 'Current Value', 'Adjustment to\nValue', 'Proposed Value', 'Indirect Cost Ratio (ICR)', 'Indirect Cost Ratio (ICR)', 'Indirect Cost Ratio (ICR)']): {
		"name":"Proposed Investment Portfolio Table",
        "process": 3,
		"processDescription": "Proposed Investment Portfolio",
	},
}


def format_investment_portfolio_recommendation_section(document):
	"""
	Modify Proposed Investment Portfolio Table Properties
	Using tuple as identifier to search table header
	"""
	#find all table from TABLE_IDENTIFIERS and save it to table list
	# tableList = multi_table_block_finder(document, TABLE_IDENTIFIERS_FIRST)
	tableList = multi_table_block_finder(document, TABLE_IDENTIFIERS_ALL)

	if tableList:
		modify_table_properties(tableList[0])
		modify_paragraph_text_color(tableList[0])
	else:
		State.notRunProgress = State.notRunProgress + "\n" + State.docname + " 3.1 modify Proposed Investment Portfolio"

def modify_table_properties(table):
	single_cell_rows = []
	for idx_r, row in enumerate(table.rows):
		if len(row.cells) > 1:
			while len(row.cells) > 4:
				delete_column_in_row(row, len(row.cells) - 1)

			if idx_r >= 2 and len(row.cells) > 0 and row.cells[0].text != "Product":
				set_table_row_height(row, 0.6)


				cellCount = 0
				for cell in row.cells:
					if cellCount == 0:
						cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
						cellCount = cellCount + 1
					else:
						cell.width = Inches(1.5)
						cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
		else:
			single_cell_rows.append(row)



def modify_paragraph_text_color(table):
	numberOfRow = len(table.rows)
	for i in range(numberOfRow):
		row = table.rows[i]
		if len(row.cells) > 1:
			basecolor = ''
			if "Total" in row.cells[0].text:
				baseRow = table.rows[i - 1] 
				baseCell = baseRow.cells[2]
				baseParagraph = baseCell.paragraphs[0]
				baseRun = baseParagraph.runs[0]
				color = baseRun.font.color.rgb
				if (str(color) == "FF0000"):
					cell = row.cells[2]
					for paragraph in cell.paragraphs:
						for run in paragraph.runs:
							run.font.color.rgb = color
	
def modify_mda_service_fee_section(document):
	success = None
	for block in iter_block_items(document):
		if isinstance(block, Paragraph):
			blockWithMDA = re.search(r"For MDA Only", block.text, re.IGNORECASE)
			
			if blockWithMDA:
				remove_prefix_from_paragraph(block, prefix="For MDA Only ")
				success = True
				break
	if not success:
				
		State.notRunProgress = State.notRunProgress + "\n" + State.docname + " 3.5 Remove and move sections (MDA)"

def move_invesment_management_fees_section(document):

	TABLE_TARGET_IDENTIFIERS = {
		tuple(['', '', 'Upfront', 'Ongoing', 'Ongoing']): {
		"name":"Product cost table",
		"processDescription": "Paste paragraph after product cost table",
		}
	}

	targetBlock = multi_table_block_finder(document, TABLE_TARGET_IDENTIFIERS)
	blockToMove = []
	inRange = False

	for block in iter_block_items(document):

		if isinstance(block, Paragraph):
			startBlock = re.search(r"Investment Management Fees", block.text, re.IGNORECASE)
			endBeforeBlock = re.search(r"Your Proposed Asset Allocation", block.text, re.IGNORECASE)
			
			if startBlock:
				inRange = True
			elif endBeforeBlock:
				inRange = False
				break
		if inRange:
			blockToMove.append(block)

	move_blocks_after(blockToMove,targetBlock[0])

	if not move_blocks_after(blockToMove,targetBlock[0]):
		State.notRunProgress = State.notRunProgress + "\n" + State.docname + " 3.5.1. Investment Management Fees and Buy/Sell Spreads"


def move_sales_and_purchases_section(document):
		
	targetBlock = ''
	blockToMove = []
	inRange = False

	for block in iter_block_items(document):

		if isinstance(block, Paragraph):
			startBlock = re.search(r"\bSales\b", block.text, re.IGNORECASE)
			endBeforeBlock = re.search(r"\bProposed portfolio(s?)\b", block.text, re.IGNORECASE)
			targetBlock = re.search(r"\bBuy/ Sell Spreads\b", block.text, re.IGNORECASE)
			if startBlock:
				inRange = True
			elif endBeforeBlock:
				inRange = False
			elif targetBlock:
				targetBlock = block
				break

		if inRange:
			blockToMove.append(block)

	if not move_blocks_after(blockToMove,targetBlock):
		State.notRunProgress = State.notRunProgress + "\n" + State.docname + " 3.5.2. Sales and Purchases"
	
def move_proposed_portfolio_section(document):
	targetBlock = ''
	blockToMove = []
	inRange = False

	for block in iter_block_items(document):

		if isinstance(block, Paragraph):
			startBlock = re.search(r"\bProposed portfolio(s?)\b", block.text, re.IGNORECASE)
			endBeforeBlock = re.search(r"\bCost comparison\b", block.text, re.IGNORECASE)
			targetBlock = re.search(r"\bInvestment Management Fees\b", block.text, re.IGNORECASE)
			if startBlock:
				inRange = True
			elif endBeforeBlock:
				inRange = False
			elif targetBlock:
				targetBlock = block
				break

		if inRange:
			blockToMove.append(block)



	if not move_blocks_after(blockToMove,targetBlock):
		State.notRunProgress = State.notRunProgress + "\n" + State.docname + " 3.5.3. Propose Portfolio/s table modification"

	target_elements = []
	blocklist = iter_block_items(document)
	for block in blocklist:
		if isinstance(block, Paragraph) and re.search(
			r"\bProposed portfolio(s?)\b", block.text, re.IGNORECASE
		):

			# include the matched paragraph itself
			target_elements.append(block)

			# consume following blocks until a table is found
			while True:
				next_block = next(blocklist, None)

				if next_block is None:
					break  # end of document

				if isinstance(next_block, Table):
					# remove paragraph immediately before the table
					if target_elements:
						target_elements.pop()
					break

				if isinstance(next_block, Paragraph):
					target_elements.append(next_block)

			break  # stop after processing the section

	for para in target_elements:
		if remove_paragraph(para) != None:
			State.notRunProgress = State.notRunProgress + "\n" + State.docname + " 3.5.3. Propose Portfolio/s Header Cleanup"